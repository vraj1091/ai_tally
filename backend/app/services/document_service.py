"""
Document Processing Service
Handles PDF, DOCX, TXT, Images (with OCR)
"""

import logging
from typing import Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

# Lazy imports
_pypdf2 = None
_docx = None
_pil = None
_pytesseract = None

def _load_pypdf2():
    global _pypdf2
    if _pypdf2 is None:
        try:
            import PyPDF2
            _pypdf2 = PyPDF2
        except ImportError:
            logger.warning("PyPDF2 not available - PDF processing disabled")
    return _pypdf2

def _load_docx():
    global _docx
    if _docx is None:
        try:
            from docx import Document as DocxDocument
            _docx = DocxDocument
        except ImportError:
            logger.warning("python-docx not available - DOCX processing disabled")
    return _docx

def _load_pil():
    global _pil
    if _pil is None:
        try:
            from PIL import Image
            _pil = Image
        except ImportError:
            logger.warning("Pillow not available - image processing disabled")
    return _pil

def _load_pytesseract():
    global _pytesseract
    if _pytesseract is None:
        try:
            import pytesseract
            _pytesseract = pytesseract
        except ImportError:
            logger.warning("pytesseract not available - OCR disabled")
    return _pytesseract


class DocumentService:
    """Service for document processing and text extraction"""
    
    @staticmethod
    def extract_pdf_text(file_path: str) -> str:
        """Extract text from PDF files"""
        PyPDF2 = _load_pypdf2()
        if PyPDF2 is None:
            logger.error("PyPDF2 not available")
            return ""
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            
            logger.info(f"Extracted {len(text)} characters from PDF: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_docx_text(file_path: str) -> str:
        """Extract text from DOCX files"""
        DocxDocument = _load_docx()
        if DocxDocument is None:
            logger.error("python-docx not available")
            return ""
        
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_image_text(file_path: str) -> str:
        """Extract text from images using OCR"""
        Image = _load_pil()
        pytesseract = _load_pytesseract()
        
        if Image is None or pytesseract is None:
            logger.error("Image OCR not available (missing PIL or pytesseract)")
            return ""
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            logger.info(f"Extracted {len(text)} characters from image: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting image {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_txt_text(file_path: str) -> str:
        """Read plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Read {len(text)} characters from TXT: {file_path}")
            return text
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            logger.info(f"Read {len(text)} characters from TXT (latin-1): {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text(file_path: str) -> Tuple[str, str]:
        """
        Extract text from any supported file type
        
        Returns:
            Tuple of (extracted_text, file_type)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        extractors = {
            '.pdf': ('pdf', DocumentService.extract_pdf_text),
            '.docx': ('docx', DocumentService.extract_docx_text),
            '.doc': ('doc', DocumentService.extract_docx_text),
            '.txt': ('txt', DocumentService.extract_txt_text),
            '.png': ('image', DocumentService.extract_image_text),
            '.jpg': ('image', DocumentService.extract_image_text),
            '.jpeg': ('image', DocumentService.extract_image_text),
            '.gif': ('image', DocumentService.extract_image_text),
            '.bmp': ('image', DocumentService.extract_image_text),
        }
        
        if extension in extractors:
            file_type, extractor = extractors[extension]
            text = extractor(file_path)
            return text, file_type
        else:
            raise ValueError(f"Unsupported file type: {extension}")
