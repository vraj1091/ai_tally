"""
DOCX Text Extractor
Advanced Word document processing
"""

from docx import Document
from typing import List, Dict
import logging

logger = logging.getLogger(__name__)

class DOCXExtractor:
    """Advanced DOCX text extraction"""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract all text from DOCX"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    @staticmethod
    def extract_paragraphs(file_path: str) -> List[str]:
        """Extract paragraphs separately"""
        try:
            doc = Document(file_path)
            return [para.text for para in doc.paragraphs if para.text.strip()]
        except Exception as e:
            logger.error(f"Error extracting paragraphs: {e}")
            raise
    
    @staticmethod
    def extract_tables(file_path: str) -> List[List[List[str]]]:
        """Extract tables from DOCX"""
        try:
            doc = Document(file_path)
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return tables
        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            raise
 
